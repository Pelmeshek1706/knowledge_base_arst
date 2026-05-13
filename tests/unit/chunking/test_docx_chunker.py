from __future__ import annotations

from pathlib import Path

from docx import Document

from personal_kb.chunking import DocxChunker
from personal_kb.parsers import DocxParser

PROJECT_ROOT = Path(__file__).resolve().parents[3]


def test_docx_chunker_preserves_heading_context_on_oversized_sections() -> None:
    file_path = PROJECT_ROOT / "data" / "DRAFT_Tpl_Application_Form_Part_B_HE_EIC_UKRAINE_2025_AIREST_5.docx"
    parsed = DocxParser().parse(
        file_path,
        source_id="data/DRAFT_Tpl_Application_Form_Part_B_HE_EIC_UKRAINE_2025_AIREST_5.docx",
    )

    chunks = DocxChunker(chunk_size=1200, chunk_overlap=100).chunk(
        parsed,
        document_id="doc-docx-1",
    )

    assert chunks[0].source_ref.section == (
        "Multimodal Digital Biomarkers for Evidence-Based Mental Health Decision Support and Monitoring"
    )
    tech_chunks = [
        chunk
        for chunk in chunks
        if chunk.source_ref.section == "Excellence > Technological breakthrough"
    ]
    assert tech_chunks
    assert all(
        chunk.text.startswith("Excellence > Technological breakthrough\n\n")
        for chunk in tech_chunks
    )


def test_docx_chunker_keeps_small_sections_intact(tmp_path: Path) -> None:
    file_path = tmp_path / "small_sections.docx"
    document = Document()
    document.add_paragraph("1. Overview")
    document.add_paragraph("Short intro paragraph.")
    document.add_paragraph("1.1 Details")
    document.add_paragraph("Small follow-up section.")
    document.save(file_path)

    parsed = DocxParser().parse(file_path, source_id="tests/fixtures/small_sections.docx")

    chunks = DocxChunker(chunk_size=500, chunk_overlap=100).chunk(
        parsed,
        document_id="doc-docx-2",
    )

    assert len(chunks) == len(parsed.structured_blocks)
    assert [chunk.source_ref.section for chunk in chunks] == ["Overview", "Overview > Details"]
